#!/usr/bin/env python3
"""
Report generation module for YMYL Audit Tool
Converts markdown reports to Word documents
"""

import io
import re
from typing import Optional
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils.helpers import safe_log

class WordReportGenerator:
    """Converts markdown reports to professionally formatted Word documents"""
    
    def __init__(self):
        """Initialize the report generator"""
        safe_log("WordReportGenerator initialized")

    def generate_report(self, markdown_content: str, title: str = "YMYL Compliance Audit Report", 
                       casino_mode: bool = False) -> bytes:
        """
        Generate Word document from markdown content
        
        Args:
            markdown_content: Markdown content to convert
            title: Document title
            casino_mode: Whether this is a casino-specific report
            
        Returns:
            Word document as bytes
        """
        try:
            safe_log(f"Generating Word report ({len(markdown_content):,} characters)")
            
            # Create document
            doc = Document()
            
            # Setup document properties
            self._setup_document(doc, title, casino_mode)
            
            # Parse and add content
            self._parse_markdown_content(doc, markdown_content)
            
            # Add footer
            self._add_footer(doc)
            
            # Save to memory
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            safe_log("Word document generation successful")
            return doc_buffer.getvalue()
            
        except Exception as e:
            safe_log(f"Word generation error: {e}")
            return self._create_error_document(str(e))

    def _setup_document(self, doc: Document, title: str, casino_mode: bool):
        """Setup document properties and styles"""
        # Document properties
        properties = doc.core_properties
        properties.title = title
        properties.author = "YMYL Audit Tool"
        properties.subject = "Casino YMYL Compliance Report" if casino_mode else "YMYL Compliance Report"
        properties.category = "Compliance Report"
        
        # Setup default styles
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(11)
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.line_spacing = 1.15

    def _parse_markdown_content(self, doc: Document, markdown_content: str):
        """Parse markdown and add to document"""
        lines = markdown_content.split('\n')
        skip_next_empty = False
        
        for i, line in enumerate(lines):
            original_line = line
            line = line.strip()
            
            # Handle empty lines
            if not line:
                if skip_next_empty:
                    skip_next_empty = False
                    continue
                if i > 0 and lines[i-1].strip():
                    doc.add_paragraph()
                continue
            
            # Parse different markdown elements
            if line.startswith('# '):
                # Main title
                paragraph = doc.add_paragraph(line[2:])
                paragraph.style = 'Title'
                skip_next_empty = True
                
            elif line.startswith('## '):
                # Section heading
                paragraph = doc.add_paragraph(line[3:])
                paragraph.style = 'Heading 1'
                skip_next_empty = True
                
            elif line.startswith('### '):
                # Subsection heading
                paragraph = doc.add_paragraph(line[4:])
                paragraph.style = 'Heading 2'
                skip_next_empty = True
                
            elif line.startswith('**') and line.endswith('**') and len(line) > 4:
                # Bold text paragraph
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
                
            elif line.startswith('- '):
                # Bullet points
                bullet_text = line[2:]
                
                if self._contains_severity_indicator(bullet_text):
                    # Special formatting for severity indicators
                    p = doc.add_paragraph(style='List Bullet')
                    self._add_severity_formatted_text(p, bullet_text)
                elif '**' in bullet_text:
                    # Handle bolded parts in bullet points
                    p = doc.add_paragraph(style='List Bullet')
                    self._add_formatted_text_to_paragraph(p, bullet_text)
                else:
                    doc.add_paragraph(bullet_text, style='List Bullet')
                    
            elif line.startswith('---'):
                # Horizontal rule
                doc.add_paragraph()
                p = doc.add_paragraph('_' * 50)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.runs[0]
                run.font.color.rgb = RGBColor(192, 192, 192)
                doc.add_paragraph()
                
            elif self._contains_severity_indicator(line):
                # Severity indicators
                p = doc.add_paragraph()
                self._add_severity_formatted_text(p, line)
                
            else:
                # Regular paragraph
                if line:
                    if '**' in line:
                        p = doc.add_paragraph()
                        self._add_formatted_text_to_paragraph(p, line)
                    else:
                        doc.add_paragraph(line)

    def _add_formatted_text_to_paragraph(self, paragraph, text):
        """Add text with embedded bold formatting"""
        parts = re.split(r'(\*\*.*?\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                # Bold text
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part:
                # Regular text
                paragraph.add_run(part)

    def _contains_severity_indicator(self, line: str) -> bool:
        """Check if line contains severity indicators"""
        return any(indicator in line for indicator in ['🔴', '🟠', '🟡', '🔵', '✅', '❌'])

    def _add_severity_formatted_text(self, paragraph, text):
        """Add severity indicator text with color formatting"""
        # Replace emoji with text for better compatibility
        emoji_replacements = {
            '🔴': '[CRITICAL]',
            '🟠': '[HIGH]',
            '🟡': '[MEDIUM]',
            '🔵': '[LOW]',
            '✅': '[✓]',
            '❌': '[✗]'
        }
        
        display_text = text
        severity_color = None
        
        for emoji, replacement in emoji_replacements.items():
            if emoji in display_text:
                display_text = display_text.replace(emoji, replacement)
                # Set color based on severity
                if emoji == '🔴':
                    severity_color = RGBColor(231, 76, 60)  # Red
                elif emoji == '🟠':
                    severity_color = RGBColor(255, 152, 0)  # Orange
                elif emoji == '🟡':
                    severity_color = RGBColor(243, 156, 18)  # Yellow
                elif emoji == '🔵':
                    severity_color = RGBColor(52, 152, 219)  # Blue
        
        # Add formatted text
        if '**' in display_text:
            self._add_formatted_text_to_paragraph(paragraph, display_text)
        else:
            run = paragraph.add_run(display_text)
        
        # Apply severity color
        if severity_color:
            for run in paragraph.runs:
                run.font.color.rgb = severity_color
                run.bold = True

    def _add_footer(self, doc: Document):
        """Add simple footer"""
        try:
            footer = doc.sections[0].footer
            footer_para = footer.paragraphs[0]
            footer_para.text = "YMYL Compliance Audit Report"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.style.font.size = Pt(9)
            footer_para.style.font.color.rgb = RGBColor(127, 140, 141)
        except Exception as e:
            safe_log(f"Could not add footer: {e}")

    def _create_error_document(self, error_message: str) -> bytes:
        """Create error document when generation fails"""
        try:
            doc = Document()
            doc.add_heading('Report Generation Error', 0)
            doc.add_paragraph(f'Failed to generate report: {error_message}')
            doc.add_paragraph('Please try again or contact support.')
            
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            return doc_buffer.getvalue()
            
        except Exception:
            return b"Error creating Word document"


# Convenience function for external use
def generate_word_report(markdown_content: str, title: str = "YMYL Compliance Audit Report", 
                        casino_mode: bool = False) -> bytes:
    """
    Generate Word document from markdown content
    
    Args:
        markdown_content: Markdown content to convert
        title: Document title
        casino_mode: Whether this is a casino-specific report
        
    Returns:
        Word document as bytes
    """
    generator = WordReportGenerator()
    return generator.generate_report(markdown_content, title, casino_mode)